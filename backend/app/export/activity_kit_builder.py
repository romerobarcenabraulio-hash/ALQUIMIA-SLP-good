"""
Plantillas, checklists y README por actividad/etapa del portafolio.
"""
from __future__ import annotations

import io
from typing import Any

from app.export.gantt_hierarchy import GanttActividad, GanttEtapa, GanttPhase


def build_activity_readme(
    actividad: GanttActividad,
    etapa: GanttEtapa,
    fase: GanttPhase,
    municipio: str,
    zm: str,
) -> str:
    lines = [
        f"ALQUIMIA · {municipio.title()} / {zm}",
        f"Fase: {fase.nombre} ({fase.phase_id})",
        f"Etapa: {etapa.nombre} ({etapa.task_id})",
        f"Actividad: {actividad.nombre}",
        "",
        f"ID actividad: {actividad.actividad_id}",
        f"Fecha sugerida: {actividad.fecha_calendario or 'Por definir'}",
        f"Semana relativa: {actividad.semana_relativa}",
        f"Responsable: {actividad.responsable}",
        "",
        "Qué hacer:",
        f"  1. Ejecutar: {actividad.nombre}",
        "  2. Registrar evidencia en bitácora de etapa",
        "  3. Marcar checklist de actividad antes de avanzar",
        "",
        "Entregables de la etapa: carpeta ../entregables/",
        "Herramientas de la etapa: carpeta ../herramientas/",
    ]
    if actividad.es_hito:
        lines.append("")
        lines.append("⚠ HITO DE CIERRE — validar gate antes de continuar.")
    return "\n".join(lines)


def build_root_readme(municipio: str, zm: str, n_fases: int, n_etapas: int) -> str:
    return f"""ALQUIMIA — Portafolio Municipal
Municipio: {municipio.title()} · ZM {zm}

ESTRUCTURA DEL PAQUETE
======================

1. analisis/
   Diagnóstico, decisión y trazabilidad. Para Presidencia, Cabildo y Tesorería.
   Empiece aquí si necesita decidir si procede el programa.

2. implementacion/
   Ejecución paso a paso según Gantt Maestro ({n_fases} fases, {n_etapas} etapas).
   Para Dirección de Servicios Públicos, Obras y operadores.

CÓMO USAR
=========
• Decisión: lea analisis/01_Resumen_Ejecutivo/ y analisis/00_Paquete_Integral_Analisis.pdf
• Operación: abra implementacion/00_Maestro/00_Guia_Implementacion.pdf
• Día a día: navegue implementacion/Fxx_.../Exx_.../actividades/Axx_.../README.txt

Generado por ÁGORA GOV · ALQUIMIA
"""


def _reportlab_pdf(title: str, body_lines: list[str]) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.pdfgen import canvas
    except ImportError:
        return b""

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - inch
    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, y, title[:80])
    y -= 0.4 * inch
    c.setFont("Helvetica", 10)
    for line in body_lines:
        if y < inch:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - inch
        c.drawString(inch, y, line[:100])
        y -= 14
    c.save()
    buf.seek(0)
    return buf.read()


def build_activity_checklist_pdf(actividad: GanttActividad, etapa: GanttEtapa) -> bytes:
    lines = [
        f"Etapa: {etapa.task_id} — {etapa.nombre[:60]}",
        f"Actividad: {actividad.nombre[:70]}",
        f"Responsable: {actividad.responsable}",
        "",
        "[ ] Actividad ejecutada según descripción",
        "[ ] Evidencia registrada (foto, acta o bitácora)",
        "[ ] Coordinación con responsable municipal confirmada",
        "",
        "Firma: _________________________  Fecha: __________",
    ]
    return _reportlab_pdf(f"Checklist — {actividad.actividad_id}", lines)


def build_etapa_checklist_pdf(etapa: GanttEtapa, entregables: list[str]) -> bytes:
    lines = [
        f"Etapa {etapa.task_id}: {etapa.nombre}",
        f"Duración: {etapa.duracion_semanas} sem · Inicio semana {etapa.inicio_semana}",
        f"Responsable: {etapa.responsable}",
        f"Predecesoras: {', '.join(etapa.predecesoras) or 'Ninguna'}",
        f"Crítica: {'Sí' if etapa.es_critica else 'No'}",
        "",
        "Entregables esperados:",
    ]
    for e in entregables or ["(ninguno específico)"]:
        lines.append(f"  • {e}")
    lines.extend([
        "",
        "[ ] Todas las actividades de la etapa completadas",
        "[ ] Entregables en carpeta entregables/ revisados",
        "[ ] Gate de fase validado",
        "",
        "Firma responsable: _________________________",
    ])
    return _reportlab_pdf(f"Checklist etapa {etapa.task_id}", lines)


def build_fase_readme_pdf(fase: GanttPhase) -> bytes:
    lines = [
        f"Fase {fase.phase_id}: {fase.nombre}",
        fase.descripcion,
        "",
        "Etapas incluidas:",
    ]
    for e in fase.etapas:
        lines.append(f"  • {e.task_id} — {e.nombre[:55]}")
    lines.extend(["", "Revise entregables/ y herramientas/ en cada etapa."])
    return _reportlab_pdf(f"Fase {fase.phase_id}", lines)


def build_etapa_readme_pdf(etapa: GanttEtapa) -> bytes:
    lines = [
        f"Etapa {etapa.task_id}: {etapa.nombre}",
        etapa.descripcion or "",
        f"Responsable: {etapa.responsable}",
        f"Inicio semana {etapa.inicio_semana} · Duración {etapa.duracion_semanas} sem",
        f"Costo estimado: ${etapa.costo_mxn:,.0f} MXN",
        f"Actividades diarias: {len(etapa.actividades)}",
    ]
    return _reportlab_pdf(f"Etapa {etapa.task_id}", lines)


def build_guia_implementacion_pdf(
    municipio: str,
    zm: str,
    hierarchy: list[GanttPhase],
) -> bytes:
    lines = [
        f"Guía de implementación — {municipio.title()} / {zm}",
        "",
        "Árbol Fase → Etapa → Actividad:",
    ]
    for fase in hierarchy:
        lines.append(f"{fase.phase_id} {fase.nombre}")
        for etapa in fase.etapas:
            lines.append(f"  {etapa.task_id} {etapa.nombre[:50]} ({len(etapa.actividades)} act.)")
    return _reportlab_pdf("Guía de implementación", lines)


def build_guia_analisis_pdf(municipio: str, zm: str) -> bytes:
    lines = [
        f"Guía de análisis — {municipio.title()} / {zm}",
        "",
        "Contenido de analisis/:",
        "  01 Resumen ejecutivo municipal (PDF + DOCX)",
        "  02 Modelo técnico-financiero + XLSX CFO",
        "  03 Diagnóstico y reforma reglamentaria",
        "  04 Coordinación metropolitana (si aplica)",
        "  07 Fuentes y trazabilidad",
        "",
        "Use 00_Paquete_Integral_Analisis.pdf para lectura completa.",
    ]
    return _reportlab_pdf("Guía de análisis", lines)


def build_pert_summary_pdf(pert_plan: Any) -> bytes:
    lines = [
        f"Ruta crítica PERT — {pert_plan.municipio} / {pert_plan.zm}",
        f"Duración total: {pert_plan.duracion_total_semanas} semanas",
        "",
        "Nodos críticos:",
    ]
    for n in pert_plan.nodes:
        if n.es_critico:
            lines.append(f"  • {n.node_id} {n.nombre[:50]} ({n.tiempo_esperado} sem)")
    return _reportlab_pdf("PERT — Ruta crítica", lines)


def build_raci_summary_pdf(raci_plan: Any) -> bytes:
    lines = [
        f"Matriz RACI — {raci_plan.municipio} / {raci_plan.zm}",
        "",
    ]
    for row in raci_plan.filas[:12]:
        lines.append(f"{row.proceso[:45]} → R: {row.responsable[:30]}")
    return _reportlab_pdf("Matriz RACI", lines)


# ─── Plantillas DOCX ─────────────────────────────────────────────────────────

_TOOL_TITLES: dict[str, str] = {
    "formulario_levantamiento_predial": "Formulario de levantamiento predial",
    "acta_sesion_cabildo": "Acta de sesión de cabildo",
    "plantilla_convocatoria_licitacion": "Plantilla de convocatoria a licitación",
    "bitacora_obra_ca": "Bitácora de obra — Centro de Acopio",
    "ficha_tecnica_vehicular": "Ficha técnica vehicular",
    "encuesta_ciudadana": "Encuesta ciudadana — línea base",
    "kpis_operativos_tracking": "Seguimiento KPIs operativos",
    "checklist_arranque_oficial": "Checklist de arranque oficial",
}


def build_tool_docx(tool_id: str, context: dict[str, str]) -> bytes:
    try:
        from docx import Document
    except ImportError:
        return b""

    municipio = context.get("municipio", "Municipio")
    zm = context.get("zm", "ZM")
    title = _TOOL_TITLES.get(tool_id, tool_id.replace("_", " ").title())

    doc = Document()
    doc.add_heading(f"{title} — {municipio.title()}", level=1)
    doc.add_paragraph(f"Zona Metropolitana: {zm}")
    doc.add_paragraph(f"Fecha: {context.get('fecha', '_______________')}")
    doc.add_paragraph("")
    doc.add_heading("Instrucciones", level=2)
    doc.add_paragraph(
        f"Complete este formato durante la etapa correspondiente del programa "
        f"ALQUIMIA en {municipio.title()}. Conserve copia firmada para trazabilidad."
    )
    doc.add_heading("Registro", level=2)
    for i in range(1, 6):
        doc.add_paragraph(f"{i}. _________________________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Responsable: _________________________")
    doc.add_paragraph("Firma: _________________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_tool_xlsx(tool_id: str, context: dict[str, str]) -> bytes:
    if tool_id != "ficha_tecnica_vehicular" and tool_id != "kpis_operativos_tracking":
        return b""
    try:
        import openpyxl
    except ImportError:
        return b""

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Datos"
    municipio = context.get("municipio", "")
    ws.append(["ALQUIMIA", municipio, context.get("zm", "")])
    if tool_id == "ficha_tecnica_vehicular":
        ws.append(["Campo", "Valor"])
        for field in ["Marca", "Modelo", "Capacidad ton", "Año", "Placa", "Costo MXN"]:
            ws.append([field, ""])
    else:
        ws.append(["KPI", "Meta", "Real", "Semana"])
        for kpi in ["Toneladas/día", "Participación %", "Ingresos MXN", "Rutas activas"]:
            ws.append([kpi, "", "", ""])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
