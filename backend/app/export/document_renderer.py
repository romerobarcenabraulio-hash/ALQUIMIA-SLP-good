"""
Fase 4 — document_renderer.py

Renderiza DraftDocument / contenido Markdown a DOCX institucional.

Reglas:
- Nunca renombrar Markdown como DOCX sin transformación real.
- No inventar cifras ni fuentes.
- Si un documento tiene is_fallback=True, DOCX lleva marca de borrador visible.
- Si ValidationReport tiene errores, DOCX lleva marca BLOQUEADO.
- ClaimLedger se incluye como anexo de trazabilidad.
- Pie de página con package_id, versión y fecha.
"""
from __future__ import annotations

import io
import logging
import re
from datetime import date
from typing import Optional

from app.export.schemas import DocumentTheme

logger = logging.getLogger(__name__)


# ─── Helpers internos ─────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))  # type: ignore[return-value]


def _add_heading(doc, text: str, level: int) -> None:
    from docx.shared import Pt, RGBColor
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.size = Pt(20)
    elif level == 2:
        p.runs[0].font.size = Pt(14)
    elif level == 3:
        p.runs[0].font.size = Pt(12)


def _add_warning_box(doc, text: str) -> None:
    """Párrafo de advertencia con fondo visual."""
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD4, 0x88, 0x1E)


def _add_kv_table(doc, rows: list[tuple[str, str]], title: str = "") -> None:
    """Tabla de 2 columnas clave-valor."""
    from docx.shared import Pt, RGBColor
    if title:
        doc.add_paragraph(title).runs[0].bold = True if title else False
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i + 1].cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph("")


def _add_footer(doc, theme: DocumentTheme, package_id: str = "") -> None:
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    run = para.add_run(
        f"{theme.footer_text}  |  ZM: {theme.zm}  |  "
        f"v{theme.version}  |  {theme.date or date.today().isoformat()}"
        + (f"  |  ID: {package_id[:8]}…" if package_id else "")
    )
    from docx.shared import Pt
    run.font.size = Pt(8)


# ─── Parser de Markdown → estructura ──────────────────────────────────────────

def _parse_markdown_blocks(md: str) -> list[tuple[str, str]]:
    """
    Parsea Markdown a lista de (tipo, contenido).
    Tipos: h1, h2, h3, quote, bullet, bold_para, para, hr, empty
    """
    blocks: list[tuple[str, str]] = []
    for line in md.splitlines():
        stripped = line.strip()
        if not stripped:
            blocks.append(("empty", ""))
        elif stripped.startswith("### "):
            blocks.append(("h3", stripped[4:]))
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            blocks.append(("h1", stripped[2:]))
        elif stripped.startswith("> "):
            blocks.append(("quote", stripped[2:]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append(("bullet", stripped[2:]))
        elif stripped.startswith("---"):
            blocks.append(("hr", ""))
        else:
            blocks.append(("para", stripped))
    return blocks


def _render_inline(para, text: str) -> None:
    """Renderiza texto inline con negrita/cursiva a un párrafo docx."""
    # Patrones: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            para.add_run(text[cursor:match.start()])
        full = match.group(0)
        if full.startswith("**"):
            r = para.add_run(match.group(2))
            r.bold = True
        elif full.startswith("`"):
            r = para.add_run(match.group(4))
            r.font.name = "Courier New"
        else:
            r = para.add_run(match.group(3))
            r.italic = True
        cursor = match.end()
    if cursor < len(text):
        para.add_run(text[cursor:])


# ─── Renderer principal ───────────────────────────────────────────────────────

def render_docx(
    md_content:        str,
    theme:             DocumentTheme,
    doc_meta:          dict,
    package_id:        str = "",
    claim_ledger_rows: Optional[list[tuple[str, str, str]]] = None,
) -> bytes:
    """
    Convierte contenido Markdown + metadata a DOCX institucional real.

    doc_meta esperado:
      document_id, titulo, audiencia, decision, status, version, source,
      is_fallback, warnings

    Retorna bytes del DOCX — NO renombra Markdown, transforma estructura real.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── Portada ───────────────────────────────────────────────────────────────
    titulo = doc_meta.get("titulo", doc_meta.get("document_id", "Documento ALQUIMIA"))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(titulo)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(*_hex_to_rgb(theme.color_primary))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run(f"{theme.brand_name}  ·  ZM {theme.zm}  ·  {theme.municipio}")

    doc.add_paragraph("")

    # ── Control de versión ────────────────────────────────────────────────────
    status = doc_meta.get("status", "borrador")
    source = doc_meta.get("source", "template")
    version = doc_meta.get("version", theme.version)

    _add_kv_table(doc, [
        ("Versión",         version),
        ("Estado",          status.upper()),
        ("Fuente",          "LLM verificado" if source == "llm" else "Template — requiere revisión"),
        ("Fecha",           theme.date or date.today().isoformat()),
        ("ZM",              theme.zm),
        ("Municipio",       theme.municipio),
        ("Package ID",      package_id[:16] + "…" if len(package_id) > 16 else package_id),
    ], title="")

    # ── Advertencias críticas ─────────────────────────────────────────────────
    is_fallback = doc_meta.get("is_fallback", False)
    is_bloqueado = (status == "bloqueado")

    if is_bloqueado:
        _add_warning_box(doc,
            "DOCUMENTO BLOQUEADO — Este borrador tiene errores críticos de validación. "
            "No puede presentarse como documento institucional hasta resolver los bloqueos."
        )

    if is_fallback and not is_bloqueado:
        _add_warning_box(doc,
            "BORRADOR AUTOMÁTICO — Generado por template, no por LLM verificado. "
            "Requiere revisión editorial antes de uso institucional."
        )

    for w in doc_meta.get("warnings", []):
        _add_warning_box(doc, w)

    # ── Audiencia y Decisión ──────────────────────────────────────────────────
    audiencia  = doc_meta.get("audiencia", [])
    decision   = doc_meta.get("decision", "")

    if audiencia or decision:
        doc.add_heading("Contexto del documento", level=2)
        if audiencia:
            p = doc.add_paragraph()
            p.add_run("Audiencia: ").bold = True
            p.add_run(", ".join(audiencia) if isinstance(audiencia, list) else str(audiencia))
        if decision:
            p = doc.add_paragraph()
            p.add_run("Decisión que habilita: ").bold = True
            p.add_run(str(decision))

    doc.add_paragraph("")

    # ── Cuerpo desde Markdown ─────────────────────────────────────────────────
    blocks = _parse_markdown_blocks(md_content)
    consecutive_empty = 0
    for kind, content in blocks:
        if kind == "empty":
            consecutive_empty += 1
            if consecutive_empty <= 1:
                doc.add_paragraph("")
            continue
        consecutive_empty = 0

        if kind == "h1":
            _add_heading(doc, content, level=1)
        elif kind == "h2":
            _add_heading(doc, content, level=2)
        elif kind == "h3":
            _add_heading(doc, content, level=3)
        elif kind == "hr":
            doc.add_paragraph("─" * 40)
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
            r = p.add_run(content)
            r.italic = True
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, content)
        else:  # para, bold_para
            p = doc.add_paragraph()
            _render_inline(p, content)

    # ── Advertencias activas al final del cuerpo ──────────────────────────────
    all_warnings = doc_meta.get("warnings", [])
    if all_warnings:
        doc.add_heading("Advertencias activas", level=2)
        for w in all_warnings:
            doc.add_paragraph(f"⚠ {w}", style="List Bullet")

    # ── Fuentes declaradas ────────────────────────────────────────────────────
    fuentes = doc_meta.get("fuentes", [])
    if fuentes:
        doc.add_heading("Fuentes y trazabilidad", level=2)
        for f in fuentes:
            doc.add_paragraph(f"• {f}", style="List Bullet")
    else:
        doc.add_heading("Fuentes y trazabilidad", level=2)
        doc.add_paragraph(
            "Las fuentes de datos usadas en este documento provienen del simulador ALQUIMIA "
            "y de los datos declarados con DataProvenance en el bundle. "
            "Ver manifest.json del paquete para detalle completo."
        )

    # ── Anexo: ClaimLedger (trazabilidad de afirmaciones) ────────────────────
    if claim_ledger_rows:
        doc.add_heading("Anexo: Trazabilidad de Afirmaciones (ClaimLedger)", level=2)
        tbl = doc.add_table(rows=len(claim_ledger_rows) + 1, cols=3)
        tbl.style = "Table Grid"
        headers = tbl.rows[0].cells
        headers[0].text = "Afirmación"
        headers[1].text = "Estado Fuente"
        headers[2].text = "Lenguaje permitido"
        for h in headers:
            for p in h.paragraphs:
                for r in p.runs:
                    r.bold = True
        for i, (claim, src_status, lang) in enumerate(claim_ledger_rows):
            row = tbl.rows[i + 1].cells
            row[0].text = claim[:120]
            row[1].text = src_status
            row[2].text = lang
        doc.add_paragraph("")

    # ── Pie de página ─────────────────────────────────────────────────────────
    _add_footer(doc, theme, package_id)

    # ── Serializar a bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
